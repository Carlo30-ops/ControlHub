import os
import sys
import hashlib
import io
from collections import defaultdict

# Agregar el directorio padre al path para importar pdf_utils
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

try:
    import fitz  # pymupdf
    try:
        import pytesseract
    except ImportError:
        pytesseract = None
except ImportError:
    pass

from pdf_utils import resolve_tesseract_path


def handle_pdf_to_word(data):
    input_file = os.path.abspath(data.get("input", ""))
    output_file = os.path.abspath(data.get("output", ""))

    # Validaciones mejoradas
    if not input_file:
        return {"ok": False, "error": "No se proporcionó archivo de entrada"}
    
    if not os.path.exists(input_file):
        return {"ok": False, "error": f"Archivo de entrada no encontrado: {input_file}"}
    
    if not output_file:
        return {"ok": False, "error": "No se proporcionó archivo de salida"}
    
    # Validar que el archivo de salida tenga extensión .docx
    if not output_file.lower().endswith('.docx'):
        return {"ok": False, "error": "El archivo de salida debe tener extensión .docx"}

    def _is_pdf_protected(path):
        try:
            doc = fitz.open(path)
            protected = doc.needs_pass
            doc.close()
            return protected
        except Exception:
            return False

    if _is_pdf_protected(input_file):
        return {
            "ok": False,
            "error": "El PDF está protegido con contraseña. Usa la herramienta 'Desbloquear PDF' primero."
        }
    
    # Verificar directorio de salida
    output_dir = os.path.dirname(output_file)
    if output_dir and not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
        except Exception as e:
            return {"ok": False, "error": f"No se pudo crear el directorio de salida: {str(e)}"}

    def _validate_pdf(path):
        """Valida que el archivo sea un PDF válido."""
        try:
            doc = fitz.open(path)
            page_count = doc.page_count
            doc.close()
            if page_count == 0:
                return False, "El PDF no tiene páginas"
            return True, None
        except Exception as e:
            return False, f"El archivo no es un PDF válido: {str(e)}"

    # Validar que el archivo de entrada sea un PDF válido
    is_valid_pdf, pdf_error = _validate_pdf(input_file)
    if not is_valid_pdf:
        return {"ok": False, "error": pdf_error}

    def _validate_docx(path):
        """Verifica que el .docx resultante tenga contenido mínimo."""
        try:
            if not os.path.exists(path):
                return False, "El archivo de salida no fue creado."
            size = os.path.getsize(path)
            if size < 1024:  # menos de 1KB es sospechoso
                return False, f"El archivo resultante es demasiado pequeño ({size} bytes). Posible conversión vacía."
            return True, None
        except Exception as e:
            return False, str(e)

    def _apply_ocr_to_pdf(input_path, output_path, tesseract_path=None, lang="spa"):
        """Aplica OCR a un PDF usando pytesseract y retorna el PDF con texto."""
        if not pytesseract:
            return None, "Librería pytesseract no instalada"
        
        tesseract_exe = resolve_tesseract_path(tesseract_path)
        if not tesseract_exe:
            return None, "No se encontró el motor Tesseract OCR. Instálalo o configura la ruta en Settings."
        
        pytesseract.pytesseract.tesseract_cmd = tesseract_exe
        
        doc = fitz.open(input_path)
        page_count = doc.page_count
        output_pdf = fitz.open()
        temp_files = []
        
        try:
            sys.stderr.write(f"OCR_PROGRESS:start|{page_count}\n")
            sys.stderr.flush()
            
            for page in doc:
                pix = page.get_pixmap(dpi=300)
                img_path = output_path + f"_temp_{page.number}.png"
                temp_files.append(img_path)
                pix.save(img_path)
                pdf_bytes = pytesseract.image_to_pdf_or_hocr(img_path, lang=lang, extension='pdf')
                temp_pdf = fitz.open("pdf", pdf_bytes)
                output_pdf.insert_pdf(temp_pdf)
                sys.stderr.write(f"OCR_PROGRESS:page|{page.number + 1}|{page_count}\n")
                sys.stderr.flush()
            
            sys.stderr.write("OCR_PROGRESS:complete\n")
            sys.stderr.flush()
        finally:
            for f in temp_files:
                if os.path.exists(f):
                    os.remove(f)
            doc.close()
        
        output_pdf.save(output_path)
        output_pdf.close()
        return output_path, None

    def _classify_pdf(path, sample_pages=5):
        """
        Analiza el contenido de un PDF y retorna su perfil.
        Solo usa fitz/PyMuPDF — sin dependencias adicionales.
        """
        sys.stderr.write("CLASSIFY_PROGRESS:start\n")
        sys.stderr.flush()
        
        try:
            doc = fitz.open(path)
            page_count = doc.page_count
            pages_to_sample = min(sample_pages, page_count)
            
            sys.stderr.write(f"CLASSIFY_PROGRESS:pages|{page_count}|{pages_to_sample}\n")
            sys.stderr.flush()
            
            text_chars = []
            image_counts = []
            drawing_counts = []
            has_forms = False
            column_layouts = 0
            has_watermark = False
            watermark_text_samples = []
            
            for i in range(pages_to_sample):
                page = doc[i]
                
                # Dimensión 1: texto
                words = page.get_text("words")
                chars = sum(len(w[4]) for w in words) if words else 0
                text_chars.append(chars)
                
                # Dimensión 2: imágenes
                images = page.get_images()
                image_counts.append(len(images))
                
                # Dimensión 3: trazos vectoriales
                drawings = page.get_drawings()
                drawing_counts.append(len(drawings))
                
                # Dimensión 4: formularios/firmas digitales
                try:
                    widgets = list(page.widgets())
                    if widgets:
                        has_forms = True
                except Exception:
                    pass
                
                # Dimensión 5: detectar layout multi-columna
                # Si los bloques de texto tienen posiciones X muy dispersas
                blocks = page.get_text("blocks")
                if blocks and len(blocks) > 4:
                    x_positions = [b[0] for b in blocks if b[6] == 0]
                    if x_positions:
                        x_range = max(x_positions) - min(x_positions)
                        page_width = page.rect.width
                        if x_range > page_width * 0.4:
                            column_layouts += 1
                
                # Dimensión 6: detectar marcas de agua
                # Buscar texto en posiciones típicas de marcas de agua (centro, diagonal)
                # o texto con baja frecuencia de aparición (sugerencia de marca de agua)
                if words:
                    # Buscar texto en el centro de la página (típico de marcas de agua)
                    page_rect = page.rect
                    center_x = page_rect.width / 2
                    center_y = page_rect.height / 2
                    center_zone_size = min(page_rect.width, page_rect.height) * 0.3
                    
                    for w in words:
                        x0, y0, x1, y1, text, block_no, block_type, direction = w
                        word_center_x = (x0 + x1) / 2
                        word_center_y = (y0 + y1) / 2
                        
                        # Verificar si está en zona central
                        if (abs(word_center_x - center_x) < center_zone_size and 
                            abs(word_center_y - center_y) < center_zone_size):
                            # Verificar si es texto corto (típico de marcas de agua)
                            if len(text) <= 20 and text.isupper():
                                has_watermark = True
                                if text not in watermark_text_samples:
                                    watermark_text_samples.append(text)
                    
                    # Detectar texto diagonal (característico de marcas de agua)
                    # Si hay bloques de texto con rotación inusual
                    try:
                        for block in blocks:
                            if block[6] == 0:  # bloque de texto
                                # fitz no proporciona rotación directa, pero podemos inferir
                                # por la posición y tamaño del bloque
                                pass
                    except Exception:
                        pass
            
            doc.close()
            
            avg_text = sum(text_chars) / len(text_chars) if text_chars else 0
            avg_images = sum(image_counts) / len(image_counts) if image_counts else 0
            avg_drawings = sum(drawing_counts) / len(drawing_counts) if drawing_counts else 0
            has_text = avg_text > 30
            has_images = avg_images > 0
            has_drawings = avg_drawings > 5
            is_multi_column = column_layouts > pages_to_sample * 0.4
            
            # Determinar perfil con consideración de marcas de agua
            if not has_text and has_images:
                profile = "scanned"
                confidence = "high"
                recommendation = "PDF escaneado detectado. Se usará pdf2docx con OCR."
            elif has_watermark:
                profile = "watermarked"
                confidence = "high"
                wm_text = ", ".join(watermark_text_samples[:3]) if watermark_text_samples else ""
                recommendation = f"Marca de agua detectada{(': ' + wm_text if wm_text else '')}. pdf2docx intentará preservar el layout original."
            elif has_forms or (has_drawings and not has_text):
                profile = "form_signature"
                confidence = "high"
                recommendation = "Formulario o firma digital detectada. pdf2docx intentará preservar el layout."
            elif has_text and has_images and avg_images > 1:
                profile = "mixed"
                confidence = "medium"
                recommendation = "PDF mixto (texto + imágenes). pdf2docx intentará preservar el layout."
            elif is_multi_column and has_text:
                profile = "complex_layout"
                confidence = "medium"
                recommendation = "Layout complejo detectado. La conversión puede variar en fidelidad."
            elif has_text and not has_images:
                profile = "digital_clean"
                confidence = "high"
                recommendation = "PDF digital limpio. Conversión de alta fidelidad esperada."
            else:
                profile = "digital_rich"
                confidence = "medium"
                recommendation = "PDF digital con contenido mixto. pdf2docx es el motor recomendado."
            
            sys.stderr.write(f"CLASSIFY_PROGRESS:complete|{profile}|{confidence}\n")
            sys.stderr.flush()
            
            return {
                "profile": profile,
                "has_text": has_text,
                "has_images": has_images,
                "has_drawings": has_drawings,
                "has_forms": has_forms,
                "has_watermark": has_watermark,
                "watermark_text": watermark_text_samples[:3] if watermark_text_samples else [],
                "avg_text_density": round(avg_text, 1),
                "avg_image_count": round(avg_images, 1),
                "avg_drawing_count": round(avg_drawings, 1),
                "page_count": page_count,
                "confidence": confidence,
                "recommendation": recommendation
            }
        except Exception as e:
            sys.stderr.write(f"CLASSIFY_PROGRESS:error|{str(e)}\n")
            sys.stderr.flush()
            return {
                "profile": "unknown",
                "has_text": False,
                "has_images": False,
                "has_drawings": False,
                "has_forms": False,
                "has_watermark": False,
                "watermark_text": [],
                "avg_text_density": 0,
                "avg_image_count": 0,
                "avg_drawing_count": 0,
                "page_count": 0,
                "confidence": "low",
                "recommendation": f"No se pudo analizar el PDF: {str(e)}"
            }

    def _extract_and_preserve_signatures(pdf_path, docx_path):
        """
        Extrae imágenes del PDF, clasifica firmas/logos y las preserva en el DOCX.
        Retorna: (preserved_count, warning)
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
        except ImportError:
            return 0, "python-docx no instalado. No se pueden preservar firmas."
        
        doc = fitz.open(pdf_path)
        page_count = doc.page_count
        docx_doc = Document(docx_path)
        
        # Extraer todas las imágenes del PDF
        images_by_page = defaultdict(list)
        image_hashes = defaultdict(list)  # hash -> [(page_num, xref, rect)]
        
        for page_num in range(page_count):
            page = doc[page_num]
            page_height = page.rect.height
            
            # Obtener todas las imágenes de la página
            image_list = page.get_images(full=True)
            
            for img in image_list:
                xref = img[0]
                base_image = doc.extract_image(xref)
                
                if base_image is None:
                    continue
                
                # Calcular hash de la imagen para detectar duplicados
                image_data = base_image["image"]
                image_hash = hashlib.md5(image_data).hexdigest()
                
                # Obtener posición de la imagen en la página
                image_rects = page.get_image_rects(xref)
                if not image_rects:
                    continue
                
                rect = image_rects[0]  # Usar el primer rectángulo
                
                # Guardar información de la imagen
                images_by_page[page_num].append({
                    "xref": xref,
                    "hash": image_hash,
                    "rect": rect,
                    "data": image_data,
                    "ext": base_image["ext"],
                    "width": base_image["width"],
                    "height": base_image["height"]
                })
                
                image_hashes[image_hash].append((page_num, xref, rect))
        
        doc.close()
        
        # Clasificar imágenes
        signatures_to_preserve = []
        logos = set()
        
        # Detectar logos (imágenes idénticas en todas las páginas)
        for image_hash, occurrences in image_hashes.items():
            if len(occurrences) == page_count:
                logos.add(image_hash)
        
        # Detectar firmas usando sistema de puntuación estructural (independiente de texto)
        # Abrir documento una sola vez fuera del bucle para evitar fuga de recursos
        doc_for_classification = fitz.open(pdf_path)
        
        for page_num, images in images_by_page.items():
            page = doc_for_classification[page_num]
            page_width = page.rect.width
            page_height = page.rect.height
            page_area = page_width * page_height
            
            for img_info in images:
                rect = img_info["rect"]
                image_hash = img_info["hash"]
                
                # Si es logo, no procesarlo como firma
                if image_hash in logos:
                    continue
                
                # Sistema de puntuación estructural
                score = 0
                
                # 1. Posición en página (tercio inferior)
                y_center = (rect.y0 + rect.y1) / 2
                if y_center > (page_height * (2/3)):
                    score += 2
                
                # 2. Tamaño relativo a la página
                rect_width = rect.x1 - rect.x0
                rect_height = rect.y1 - rect.y0
                rect_area = rect_width * rect_height
                width_ratio = (rect_width / page_width) * 100
                area_ratio = (rect_area / page_area) * 100
                
                if width_ratio <= 30 and area_ratio <= 5:
                    score += 2
                
                # 3. Proporción de aspecto (alargadas horizontalmente)
                aspect_ratio = img_info["width"] / img_info["height"] if img_info["height"] > 0 else 0
                if 1.8 <= aspect_ratio <= 4.0:
                    score += 2
                
                # 4. Densidad de píxeles no blancos (firmas tienen baja densidad)
                try:
                    from PIL import Image
                    pil_image = Image.open(io.BytesIO(img_info["data"]))
                    if pil_image.mode != 'RGB':
                        pil_image = pil_image.convert('RGB')
                    pixels = pil_image.getdata()
                    non_white_count = sum(1 for p in pixels if p != (255, 255, 255))
                    total_pixels = len(pixels)
                    density = (non_white_count / total_pixels) * 100 if total_pixels > 0 else 0
                    if density <= 40:
                        score += 3
                except Exception:
                    # Si falla el cálculo de densidad, no sumar puntos
                    pass
                
                # 5. No es logo (ya verificado arriba)
                score += 2
                
                # Umbral: score >= 7 para considerar como firma
                if score >= 7:
                    signatures_to_preserve.append({
                        "page_num": page_num,
                        "image_data": img_info["data"],
                        "ext": img_info["ext"],
                        "rect": rect,
                        "width": img_info["width"],
                        "height": img_info["height"],
                        "score": score
                    })
        
        doc_for_classification.close()
        
        # Preservar firmas en el DOCX (insertar al final agrupadas por página)
        preserved_count = 0
        
        if signatures_to_preserve:
            # Agrupar firmas por página
            signatures_by_page = defaultdict(list)
            for sig in signatures_to_preserve:
                signatures_by_page[sig["page_num"]].append(sig)
            
            # Agregar separador antes de las firmas
            docx_doc.add_paragraph("─" * 50)
            para = docx_doc.add_paragraph("FIRMAS PRESERVADAS DEL PDF ORIGINAL")
            para.alignment = 1  # Centrado
            para.runs[0].bold = True
            para.runs[0].font.size = Pt(12)
            
            # Insertar firmas agrupadas por página
            for page_num in sorted(signatures_by_page.keys()):
                page_sigs = signatures_by_page[page_num]
                
                # Título de página
                para = docx_doc.add_paragraph(f"\nFirmas de página {page_num + 1}:")
                para.runs[0].bold = True
                para.runs[0].font.size = Pt(10)
                
                # Insertar cada firma de esta página como imagen inline
                for sig_info in page_sigs:
                    try:
                        # Calcular tamaño en pulgadas (máximo 2 pulgadas de ancho)
                        max_width_inches = 2.0
                        aspect_ratio = sig_info["width"] / sig_info["height"]
                        
                        if sig_info["width"] > sig_info["height"]:
                            width = Inches(max_width_inches)
                            height = Inches(max_width_inches / aspect_ratio)
                        else:
                            height = Inches(max_width_inches)
                            width = Inches(max_width_inches * aspect_ratio)
                        
                        # Agregar párrafo con la imagen inline
                        para = docx_doc.add_paragraph()
                        para.alignment = 2  # Centrado
                        
                        image_stream = io.BytesIO(sig_info["image_data"])
                        run = para.add_run()
                        run.add_picture(image_stream, width=width, height=height)
                        
                        # Agregar texto descriptivo
                        para.add_run(f"\n[Firma de página {page_num + 1}]")
                        para.runs[-1].font.size = Pt(8)
                        para.runs[-1].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                        
                        preserved_count += 1
                    except Exception as e:
                        sys.stderr.write(f"SIGNATURE_PRESERVE:error|{page_num}|{str(e)}\n")
                        sys.stderr.flush()
        
        if preserved_count > 0:
            docx_doc.save(docx_path)
        
        return preserved_count, None

    pdf_info = _classify_pdf(input_file)
    profile = pdf_info["profile"]
    page_count = pdf_info["page_count"]

    def _ok_result(output_path, warning=None, pdf_profile=None, strategy_used=None, input_size=None, output_size=None):
        valid, reason = _validate_docx(output_path)
        if not valid:
            return {"ok": False, "error": f"Conversión completada pero resultado inválido: {reason}"}
        result = {
            "ok": True, 
            "output": output_path,
            "page_count": page_count,
            "pdf_profile": pdf_profile
        }
        if warning:
            result["warning"] = warning
        if strategy_used:
            result["strategy_used"] = strategy_used
        if input_size:
            result["input_size"] = input_size
        if output_size:
            result["output_size"] = output_size
        return result

    # Reportar inicio de conversión
    sys.stderr.write(f"PDF2WORD_PROGRESS:starting|{profile}|{page_count}\n")
    sys.stderr.flush()

    # Usar solo pdf2docx para evitar diálogos de Word y mejorar calidad con OCR
    strategy_order = ["pdf2docx"]

    base_warning = None
    if pdf_info["confidence"] != "high" or profile in ("mixed", "complex_layout"):
        base_warning = pdf_info["recommendation"]

    errors = []
    input_size = os.path.getsize(input_file) if os.path.exists(input_file) else 0
    
    try:
        for strategy_idx, strategy in enumerate(strategy_order):
            try:
                # Reportar estrategia actual
                sys.stderr.write(f"PDF2WORD_PROGRESS:strategy|{strategy}|{strategy_idx + 1}/{len(strategy_order)}\n")
                sys.stderr.flush()
                
                if strategy == "pdf2docx":
                    try:
                        from pdf2docx import Converter
                    except ImportError:
                        errors.append("pdf2docx: Librería no instalada. Ejecuta: pip install pdf2docx")
                        continue
                    
                    # Si el perfil es scanned, aplicar OCR primero
                    pdf_to_convert = input_file
                    ocr_warning = None
                    if profile == "scanned":
                        sys.stderr.write("PDF2WORD_PROGRESS:ocr_start\n")
                        sys.stderr.flush()
                        
                        ocr_pdf_path = input_file + "_ocr.pdf"
                        ocr_result, ocr_error = _apply_ocr_to_pdf(input_file, ocr_pdf_path, data.get("tesseract_path"), "spa")
                        
                        if ocr_error:
                            ocr_warning = f"OCR falló: {ocr_error}. Intentando conversión sin OCR."
                            pdf_to_convert = input_file
                        else:
                            pdf_to_convert = ocr_pdf_path
                            sys.stderr.write("PDF2WORD_PROGRESS:ocr_complete\n")
                            sys.stderr.flush()
                    
                    sys.stderr.write("PDF2WORD_PROGRESS:converting\n")
                    sys.stderr.flush()
                    
                    cv = Converter(pdf_to_convert)
                    
                    # Configuración mejorada según el perfil del PDF
                    # Nota: pdf2docx usa parámetros estándar, no enable_table/image/formula
                    # La detección de tablas e imágenes es automática en pdf2docx
                    if profile == "scanned":
                        # Para PDFs escaneados, priorizar OCR y preservación de imágenes
                        cv.convert(output_file,
                            start=0, end=None,
                            multi_processing=False,
                            connected_border_tolerance=0.3,  # más tolerante para OCR
                            line_overlap_threshold=0.8,
                            image_overlap_threshold=0.5,  # preservar imágenes pequeñas (firmas)
                            float_image_ignorable_gap=20.0,  # permitir superposición mayor para firmas
                            debug=False
                        )
                    elif profile == "mixed" or profile == "watermarked":
                        # Para PDFs mixtos o con marcas de agua, preservar layout
                        cv.convert(output_file,
                            start=0, end=None,
                            multi_processing=False,
                            connected_border_tolerance=0.2,
                            line_overlap_threshold=0.9,
                            image_overlap_threshold=0.5,  # preservar imágenes pequeñas (firmas)
                            float_image_ignorable_gap=20.0,  # permitir superposición mayor para firmas
                            debug=False
                        )
                    else:
                        # Configuración estándar
                        cv.convert(output_file,
                            start=0, end=None,
                            multi_processing=False,
                            connected_border_tolerance=0,
                            line_overlap_threshold=0.9,
                            image_overlap_threshold=0.5,  # preservar imágenes pequeñas (firmas)
                            float_image_ignorable_gap=20.0,  # permitir superposición mayor para firmas
                            debug=False
                        )
                    cv.close()
                    
                    # Limpiar archivo temporal de OCR si existe
                    if pdf_to_convert != input_file and os.path.exists(pdf_to_convert):
                        os.remove(pdf_to_convert)
                    
                    # Preservar firmas para perfiles digitales
                    if profile in ("digital_clean", "digital_rich", "mixed", "form_signature"):
                        sys.stderr.write("SIGNATURE_PRESERVE:start\n")
                        sys.stderr.flush()
                        
                        preserved_count, sig_warning = _extract_and_preserve_signatures(input_file, output_file)
                        
                        if sig_warning:
                            sys.stderr.write(f"SIGNATURE_PRESERVE:warning|{sig_warning}\n")
                            sys.stderr.flush()
                        elif preserved_count > 0:
                            sys.stderr.write(f"SIGNATURE_PRESERVE:complete|{preserved_count}\n")
                            sys.stderr.flush()
                    
                    sys.stderr.write("PDF2WORD_PROGRESS:complete\n")
                    sys.stderr.flush()
                    
                    output_size = os.path.getsize(output_file) if os.path.exists(output_file) else 0
                    final_warning = ocr_warning or base_warning or "Conversión via pdf2docx. Layouts complejos pueden variar."
                    return _ok_result(output_file, final_warning, profile, "pdf2docx", input_size, output_size)
            except Exception as e:
                errors.append(f"{strategy}: {str(e)}")
                continue

        return {"ok": False, "error": " | ".join(errors)}
    except Exception as e:
        return {"ok": False, "error": f"Error en conversión PDF a Word: {str(e)}"}

